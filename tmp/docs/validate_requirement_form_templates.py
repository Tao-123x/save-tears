from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT_DIR / "output" / "doc"

EXPECTED_DOCS = [
    {
        "filename": "水处理组需求采集表.docx",
        "title": "水处理组需求采集表",
        "purpose_fragment": "Purpose / 填写目的: 本表用于收集贵组对系统功能、数据对接和协作方式的需求，以便后续统一复核与开发安排。",
        "instructions_fragment": "Instructions / 填写说明: 每组填写一份；Part A 填写一次；Part B 可填写多条需求；若细节暂未确定，可先填写初步想法，后续会议再统一确认。",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field 字段",
            "Response 填写内容",
            "Group name 小组名称",
            "Main responsibility 主要负责内容",
            "Current progress or existing outputs 当前已有成果",
            "What support is needed from our system 希望我们的系统提供哪些支持",
            "What your group can provide 你们组可以提供什么",
            "Related devices, data, documents, or interfaces 相关设备、数据、文档或接口",
            "Main contact person 主要联系人",
            "Preferred collaboration method 希望的对接方式",
            "Requirement title 需求名称",
            "Problem to solve 想解决的问题",
            "Expected function 希望我们实现的功能",
            "Input needed 需要的输入",
            "Who provides the input 输入由谁提供",
            "Expected output 期望输出",
            "Target users 使用对象",
            "Priority 优先级",
            "Preferred timeline 希望完成时间",
            "Dependencies or risks 依赖项或风险",
            "Remarks 备注",
        ],
        "hint_fragments": [
            "请重点说明你们关注的水质指标、处理阶段、阈值告警和历史趋势需求。",
            "请写明你们组可以提供哪些检测数据、设备数据或文档资料。",
            "如果希望系统支持前后处理结果对比，请说明对比维度和展示方式。",
        ],
    },
    {
        "filename": "排水组需求采集表.docx",
        "title": "排水组需求采集表",
        "purpose_fragment": "Purpose / 填写目的: 本表用于收集贵组对系统功能、数据对接和协作方式的需求，以便后续统一复核与开发安排。",
        "instructions_fragment": "Instructions / 填写说明: 每组填写一份；Part A 填写一次；Part B 可填写多条需求；若细节暂未确定，可先填写初步想法，后续会议再统一确认。",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field 字段",
            "Response 填写内容",
            "Group name 小组名称",
            "Main responsibility 主要负责内容",
            "Current progress or existing outputs 当前已有成果",
            "What support is needed from our system 希望我们的系统提供哪些支持",
            "What your group can provide 你们组可以提供什么",
            "Related devices, data, documents, or interfaces 相关设备、数据、文档或接口",
            "Main contact person 主要联系人",
            "Preferred collaboration method 希望的对接方式",
            "Requirement title 需求名称",
            "Problem to solve 想解决的问题",
            "Expected function 希望我们实现的功能",
            "Input needed 需要的输入",
            "Who provides the input 输入由谁提供",
            "Expected output 期望输出",
            "Target users 使用对象",
            "Priority 优先级",
            "Preferred timeline 希望完成时间",
            "Dependencies or risks 依赖项或风险",
            "Remarks 备注",
        ],
        "hint_fragments": [
            "请重点说明需要监测的排水点、异常工况和需要追踪的记录类型。",
            "请写明是否需要实时状态、周期记录、路径展示或区域影响信息。",
            "请说明你们组能够提供哪些现场数据、设备信号或流程资料。",
        ],
    },
    {
        "filename": "应用组需求采集表.docx",
        "title": "应用组需求采集表",
        "purpose_fragment": "Purpose / 填写目的: 本表用于收集贵组对系统功能、数据对接和协作方式的需求，以便后续统一复核与开发安排。",
        "instructions_fragment": "Instructions / 填写说明: 每组填写一份；Part A 填写一次；Part B 可填写多条需求；若细节暂未确定，可先填写初步想法，后续会议再统一确认。",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field 字段",
            "Response 填写内容",
            "Group name 小组名称",
            "Main responsibility 主要负责内容",
            "Current progress or existing outputs 当前已有成果",
            "What support is needed from our system 希望我们的系统提供哪些支持",
            "What your group can provide 你们组可以提供什么",
            "Related devices, data, documents, or interfaces 相关设备、数据、文档或接口",
            "Main contact person 主要联系人",
            "Preferred collaboration method 希望的对接方式",
            "Requirement title 需求名称",
            "Problem to solve 想解决的问题",
            "Expected function 希望我们实现的功能",
            "Input needed 需要的输入",
            "Who provides the input 输入由谁提供",
            "Expected output 期望输出",
            "Target users 使用对象",
            "Priority 优先级",
            "Preferred timeline 希望完成时间",
            "Dependencies or risks 依赖项或风险",
            "Remarks 备注",
        ],
        "hint_fragments": [
            "请重点说明目标用户是谁、他们最常做什么操作、最先需要看到什么信息。",
            "请说明哪些页面、流程或交互体验最值得优先完善。",
            "请尽量写出当前系统在使用体验上的问题和你们的改进建议。",
        ],
    },
    {
        "filename": "Derection组需求采集表.docx",
        "title": "Derection组需求采集表",
        "purpose_fragment": "Purpose / 填写目的: 本表用于收集贵组对系统功能、数据对接和协作方式的需求，以便后续统一复核与开发安排。",
        "instructions_fragment": "Instructions / 填写说明: 每组填写一份；Part A 填写一次；Part B 可填写多条需求；若细节暂未确定，可先填写初步想法，后续会议再统一确认。",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field 字段",
            "Response 填写内容",
            "Group name 小组名称",
            "Main responsibility 主要负责内容",
            "Current progress or existing outputs 当前已有成果",
            "What support is needed from our system 希望我们的系统提供哪些支持",
            "What your group can provide 你们组可以提供什么",
            "Related devices, data, documents, or interfaces 相关设备、数据、文档或接口",
            "Main contact person 主要联系人",
            "Preferred collaboration method 希望的对接方式",
            "Requirement title 需求名称",
            "Problem to solve 想解决的问题",
            "Expected function 希望我们实现的功能",
            "Input needed 需要的输入",
            "Who provides the input 输入由谁提供",
            "Expected output 期望输出",
            "Target users 使用对象",
            "Priority 优先级",
            "Preferred timeline 希望完成时间",
            "Dependencies or risks 依赖项或风险",
            "Remarks 备注",
        ],
        "hint_fragments": [
            "请重点说明决策、统筹、汇报或管理时最需要看到哪些信息。",
            "请说明是否需要总览面板、阶段汇总、统计报表或会议展示材料。",
            "请尽量明确你们最关注的关键指标、更新频率和结果形式。",
        ],
    },
    {
        "filename": "Water_Treatment_Group_Requirement_Form.docx",
        "title": "Water Treatment Group Requirement Collection Form",
        "purpose_fragment": "Purpose / 填写目的: This form is used to collect your group's feature requests, technical coordination needs, and collaboration expectations for later review and planning.",
        "instructions_fragment": "Instructions / 填写说明: Each group should complete one form. Fill Part A once. Use Part B for multiple requirement items. Rough answers are acceptable if details are not final yet.",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field 字段",
            "Response 填写内容",
            "Group name 小组名称",
            "Main responsibility 主要负责内容",
            "Current progress or existing outputs 当前已有成果",
            "What support is needed from our system 希望我们的系统提供哪些支持",
            "What your group can provide 你们组可以提供什么",
            "Related devices, data, documents, or interfaces 相关设备、数据、文档或接口",
            "Main contact person 主要联系人",
            "Preferred collaboration method 希望的对接方式",
            "Requirement title 需求名称",
            "Problem to solve 想解决的问题",
            "Expected function 希望我们实现的功能",
            "Input needed 需要的输入",
            "Who provides the input 输入由谁提供",
            "Expected output 期望输出",
            "Target users 使用对象",
            "Priority 优先级",
            "Preferred timeline 希望完成时间",
            "Dependencies or risks 依赖项或风险",
            "Remarks 备注",
        ],
        "hint_fragments": [
            "Please focus on water quality indicators, treatment stages, threshold alerts, and trend-view needs.",
            "Please specify what monitoring data, device data, or documents your group can provide.",
            "If you want before-and-after treatment comparison, describe the comparison dimensions and preferred presentation format.",
        ],
    },
    {
        "filename": "Drainage_Group_Requirement_Form.docx",
        "title": "Drainage Group Requirement Collection Form",
        "purpose_fragment": "Purpose / 填写目的: This form is used to collect your group's feature requests, technical coordination needs, and collaboration expectations for later review and planning.",
        "instructions_fragment": "Instructions / 填写说明: Each group should complete one form. Fill Part A once. Use Part B for multiple requirement items. Rough answers are acceptable if details are not final yet.",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field / 字段",
            "Response / 填写内容",
            "Group name / 小组名称",
            "Main responsibility / 主要负责内容",
            "Current progress or existing outputs / 当前已有成果",
            "What support is needed from our system / 希望我们的系统提供哪些支持",
            "What your group can provide / 你们组可以提供什么",
            "Related devices, data, documents, or interfaces / 相关设备、数据、文档或接口",
            "Main contact person / 主要联系人",
            "Preferred collaboration method / 希望的对接方式",
            "Requirement title / 需求名称",
            "Problem to solve / 想解决的问题",
            "Expected function / 希望我们实现的功能",
            "Input needed / 需要的输入",
            "Who provides the input / 输入由谁提供",
            "Expected output / 期望输出",
            "Target users / 使用对象",
            "Priority / 优先级",
            "Preferred timeline / 希望完成时间",
            "Dependencies or risks / 依赖项或风险",
            "Remarks / 备注",
        ],
        "hint_fragments": [
            "Please focus on drainage points to monitor, abnormal conditions, and the record types that matter most.",
            "Please state whether you need real-time status, periodic records, route visibility, or area impact information.",
            "Please describe the field data, device signals, or process materials your group can provide.",
        ],
    },
    {
        "filename": "Application_Group_Requirement_Form.docx",
        "title": "Application Group Requirement Collection Form",
        "purpose_fragment": "Purpose / 填写目的: This form is used to collect your group's feature requests, technical coordination needs, and collaboration expectations for later review and planning.",
        "instructions_fragment": "Instructions / 填写说明: Each group should complete one form. Fill Part A once. Use Part B for multiple requirement items. Rough answers are acceptable if details are not final yet.",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field / 字段",
            "Response / 填写内容",
            "Group name / 小组名称",
            "Main responsibility / 主要负责内容",
            "Current progress or existing outputs / 当前已有成果",
            "What support is needed from our system / 希望我们的系统提供哪些支持",
            "What your group can provide / 你们组可以提供什么",
            "Related devices, data, documents, or interfaces / 相关设备、数据、文档或接口",
            "Main contact person / 主要联系人",
            "Preferred collaboration method / 希望的对接方式",
            "Requirement title / 需求名称",
            "Problem to solve / 想解决的问题",
            "Expected function / 希望我们实现的功能",
            "Input needed / 需要的输入",
            "Who provides the input / 输入由谁提供",
            "Expected output / 期望输出",
            "Target users / 使用对象",
            "Priority / 优先级",
            "Preferred timeline / 希望完成时间",
            "Dependencies or risks / 依赖项或风险",
            "Remarks / 备注",
        ],
        "hint_fragments": [
            "Please focus on target users, their key actions, and the information they should see first.",
            "Please explain which pages, flows, or interaction problems should be improved first.",
            "Please describe the most important usability issues in the current system and your suggestions.",
        ],
    },
    {
        "filename": "Derection_Group_Requirement_Form.docx",
        "title": "Derection Group Requirement Collection Form",
        "purpose_fragment": "Purpose / 填写目的: This form is used to collect your group's feature requests, technical coordination needs, and collaboration expectations for later review and planning.",
        "instructions_fragment": "Instructions / 填写说明: Each group should complete one form. Fill Part A once. Use Part B for multiple requirement items. Rough answers are acceptable if details are not final yet.",
        "section_fragments": [
            "Part A. Group Overview / 小组概况",
            "Part B. Requirement Items / 需求明细",
            "Group-specific hints / 本组填写提示",
            "Follow-up / 后续说明:",
        ],
        "table_fragments": [
            "Field / 字段",
            "Response / 填写内容",
            "Group name / 小组名称",
            "Main responsibility / 主要负责内容",
            "Current progress or existing outputs / 当前已有成果",
            "What support is needed from our system / 希望我们的系统提供哪些支持",
            "What your group can provide / 你们组可以提供什么",
            "Related devices, data, documents, or interfaces / 相关设备、数据、文档或接口",
            "Main contact person / 主要联系人",
            "Preferred collaboration method / 希望的对接方式",
            "Requirement title / 需求名称",
            "Problem to solve / 想解决的问题",
            "Expected function / 希望我们实现的功能",
            "Input needed / 需要的输入",
            "Who provides the input / 输入由谁提供",
            "Expected output / 期望输出",
            "Target users / 使用对象",
            "Priority / 优先级",
            "Preferred timeline / 希望完成时间",
            "Dependencies or risks / 依赖项或风险",
            "Remarks / 备注",
        ],
        "hint_fragments": [
            "Please focus on the information most needed for decision-making, coordination, reporting, or management.",
            "Please explain whether you need overview dashboards, stage summaries, statistics reports, or presentation materials.",
            "Please clarify the key indicators, update frequency, and preferred output format.",
        ],
    },
]


def compact_text(value):
    return "".join(ch for ch in value if not ch.isspace() and ch != "/")


def iter_text_units(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    yield text
                    yield " ".join(line.strip() for line in text.splitlines() if line.strip())


def validate_doc(path, expected):
    if not path.exists():
        raise FileNotFoundError(f"Missing expected docx file: {path}")

    doc = Document(str(path))
    content_units = [compact_text(unit) for unit in iter_text_units(doc)]

    required_fragments = [expected["title"], expected["purpose_fragment"], expected["instructions_fragment"]]
    required_fragments.extend(expected["section_fragments"])
    required_fragments.extend(expected["table_fragments"])
    required_fragments.extend(expected["hint_fragments"])

    missing = [
        fragment
        for fragment in required_fragments
        if not any(compact_text(fragment) in unit for unit in content_units)
    ]
    if missing:
        raise ValueError(f"Missing expected fragments in {path.name}: {missing!r}")


def main():
    expected_files = [OUTPUT_DIR / doc["filename"] for doc in EXPECTED_DOCS]
    for path in expected_files:
        if not path.exists():
            raise FileNotFoundError(f"Missing expected file: {path}")

    for expected in EXPECTED_DOCS:
        validate_doc(OUTPUT_DIR / expected["filename"], expected)

    print("PASS: requirement form templates are complete")


if __name__ == "__main__":
    main()
