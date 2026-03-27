import os
import shutil
import subprocess
import sys
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT_DIR / "output" / "doc"
CHINESE_FONT_CANDIDATES = [
    "PingFang SC",
    "Hiragino Sans GB",
    "Songti SC",
    "Noto Sans CJK SC",
    "Microsoft YaHei",
    "SimSun",
]


def _normalize_font_name(value):
    return "".join(ch for ch in value.lower() if ch.isalnum())


def _font_dirs():
    if sys.platform == "darwin":
        return [
            Path("/System/Library/Fonts"),
            Path("/Library/Fonts"),
            Path.home() / "Library" / "Fonts",
        ]
    if os.name == "nt":
        windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
        return [windir / "Fonts"]
    return [
        Path("/usr/share/fonts"),
        Path("/usr/local/share/fonts"),
        Path.home() / ".fonts",
        Path.home() / ".local" / "share" / "fonts",
    ]


def _match_font_from_filesystem(candidate):
    wanted = _normalize_font_name(candidate)
    for font_dir in _font_dirs():
        if not font_dir.exists():
            continue
        for pattern in ("*.ttf", "*.ttc", "*.otf", "*.dfont"):
            for path in font_dir.glob(pattern):
                stem = _normalize_font_name(path.stem)
                if wanted in stem or stem in wanted:
                    return candidate
    return None


def _match_font_with_fc_match(candidate):
    fc_match = shutil.which("fc-match")
    if not fc_match:
        return None
    try:
        result = subprocess.run(
            [fc_match, "-f", "%{family[0]}\n", candidate],
            check=True,
            capture_output=True,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return None
    matched = result.stdout.strip()
    if not matched:
        return None
    wanted = _normalize_font_name(candidate)
    got = _normalize_font_name(matched)
    if wanted in got or got in wanted:
        return candidate
    return None


@lru_cache(maxsize=1)
def resolve_chinese_font():
    for candidate in CHINESE_FONT_CANDIDATES:
        if _match_font_with_fc_match(candidate) or _match_font_from_filesystem(candidate):
            return candidate
    return CHINESE_FONT_CANDIDATES[0]


CHINESE_FONT = resolve_chinese_font()


GROUPS = [
    {
        "key": "water_treatment",
        "zh_filename": "水处理组需求采集表.docx",
        "en_filename": "Water_Treatment_Group_Requirement_Form.docx",
        "zh_title": "水处理组需求采集表",
        "en_title": "Water Treatment Group Requirement Collection Form",
        "zh_hints": [
            "请重点说明你们关注的水质指标、处理阶段、阈值告警和历史趋势需求。",
            "请写明你们组可以提供哪些检测数据、设备数据或文档资料。",
            "如果希望系统支持前后处理结果对比，请说明对比维度和展示方式。",
        ],
        "en_hints": [
            "Please focus on water quality indicators, treatment stages, threshold alerts, and trend-view needs.",
            "Please specify what monitoring data, device data, or documents your group can provide.",
            "If you want before-and-after treatment comparison, describe the comparison dimensions and preferred presentation format.",
        ],
    },
    {
        "key": "drainage",
        "zh_filename": "排水组需求采集表.docx",
        "en_filename": "Drainage_Group_Requirement_Form.docx",
        "zh_title": "排水组需求采集表",
        "en_title": "Drainage Group Requirement Collection Form",
        "zh_hints": [
            "请重点说明需要监测的排水点、异常工况和需要追踪的记录类型。",
            "请写明是否需要实时状态、周期记录、路径展示或区域影响信息。",
            "请说明你们组能够提供哪些现场数据、设备信号或流程资料。",
        ],
        "en_hints": [
            "Please focus on drainage points to monitor, abnormal conditions, and the record types that matter most.",
            "Please state whether you need real-time status, periodic records, route visibility, or area impact information.",
            "Please describe the field data, device signals, or process materials your group can provide.",
        ],
    },
    {
        "key": "application",
        "zh_filename": "应用组需求采集表.docx",
        "en_filename": "Application_Group_Requirement_Form.docx",
        "zh_title": "应用组需求采集表",
        "en_title": "Application Group Requirement Collection Form",
        "zh_hints": [
            "请重点说明目标用户是谁、他们最常做什么操作、最先需要看到什么信息。",
            "请说明哪些页面、流程或交互体验最值得优先完善。",
            "请尽量写出当前系统在使用体验上的问题和你们的改进建议。",
        ],
        "en_hints": [
            "Please focus on target users, their key actions, and the information they should see first.",
            "Please explain which pages, flows, or interaction problems should be improved first.",
            "Please describe the most important usability issues in the current system and your suggestions.",
        ],
    },
    {
        "key": "derection",
        "zh_filename": "Derection组需求采集表.docx",
        "en_filename": "Derection_Group_Requirement_Form.docx",
        "zh_title": "Derection组需求采集表",
        "en_title": "Derection Group Requirement Collection Form",
        "zh_hints": [
            "请重点说明决策、统筹、汇报或管理时最需要看到哪些信息。",
            "请说明是否需要总览面板、阶段汇总、统计报表或会议展示材料。",
            "请尽量明确你们最关注的关键指标、更新频率和结果形式。",
        ],
        "en_hints": [
            "Please focus on the information most needed for decision-making, coordination, reporting, or management.",
            "Please explain whether you need overview dashboards, stage summaries, statistics reports, or presentation materials.",
            "Please clarify the key indicators, update frequency, and preferred output format.",
        ],
    },
]


PART_A_FIELDS = [
    ("Group name", "小组名称"),
    ("Main responsibility", "主要负责内容"),
    ("Current progress or existing outputs", "当前已有成果"),
    ("What support is needed from our system", "希望我们的系统提供哪些支持"),
    ("What your group can provide", "你们组可以提供什么"),
    ("Related devices, data, documents, or interfaces", "相关设备、数据、文档或接口"),
    ("Main contact person", "主要联系人"),
    ("Preferred collaboration method", "希望的对接方式"),
]


PART_B_HEADERS = [
    ("Requirement title", "需求名称"),
    ("Problem to solve", "想解决的问题"),
    ("Expected function", "希望我们实现的功能"),
    ("Input needed", "需要的输入"),
    ("Who provides the input", "输入由谁提供"),
    ("Expected output", "期望输出"),
    ("Target users", "使用对象"),
    ("Priority", "优先级"),
    ("Preferred timeline", "希望完成时间"),
    ("Dependencies or risks", "依赖项或风险"),
    ("Remarks", "备注"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_page_setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.font.size = Pt(9.5)


def set_east_asia_font(run, font_name):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(31, 78, 121)
    set_east_asia_font(r, CHINESE_FONT)


def add_intro(doc, purpose, instructions):
    p = doc.add_paragraph()
    r = p.add_run("Purpose / 填写目的: ")
    r.bold = True
    set_east_asia_font(r, CHINESE_FONT)
    r = p.add_run(purpose)
    set_east_asia_font(r, CHINESE_FONT)

    p = doc.add_paragraph()
    r = p.add_run("Instructions / 填写说明: ")
    r.bold = True
    set_east_asia_font(r, CHINESE_FONT)
    r = p.add_run(instructions)
    set_east_asia_font(r, CHINESE_FONT)


def add_part_a_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Field / 字段"
    headers[1].text = "Response / 填写内容"
    set_cell_shading(headers[0], "D9EAF7")
    set_cell_shading(headers[1], "D9EAF7")

    for en_label, zh_label in PART_A_FIELDS:
        row = table.add_row().cells
        row[0].text = f"{en_label} / {zh_label}"
        row[1].text = ""


def add_part_b_table(doc):
    table = doc.add_table(rows=4, cols=len(PART_B_HEADERS))
    table.style = "Table Grid"
    for index, (en_label, zh_label) in enumerate(PART_B_HEADERS):
        table.rows[0].cells[index].text = f"{en_label}\n{zh_label}"
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row_index in range(1, 4):
        for col_index in range(len(PART_B_HEADERS)):
            table.rows[row_index].cells[col_index].text = ""


def add_hint_list(doc, hints, heading):
    doc.add_paragraph(heading)
    for hint in hints:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(hint)


def build_document(language, group):
    doc = Document()
    apply_page_setup(doc)

    if language == "zh":
        title = group["zh_title"]
        purpose = "本表用于收集贵组对系统功能、数据对接和协作方式的需求，以便后续统一复核与开发安排。"
        instructions = "每组填写一份；Part A 填写一次；Part B 可填写多条需求；若细节暂未确定，可先填写初步想法，后续会议再统一确认。"
        hint_heading = "Group-specific hints / 本组填写提示"
        follow_up = "Follow-up / 后续说明: 如有尚不明确的事项，可在后续统一复核会议中继续讨论并确认。"
    else:
        title = group["en_title"]
        purpose = "This form is used to collect your group's feature requests, technical coordination needs, and collaboration expectations for later review and planning."
        instructions = "Each group should complete one form. Fill Part A once. Use Part B for multiple requirement items. Rough answers are acceptable if details are not final yet."
        hint_heading = "Group-specific hints / 本组填写提示"
        follow_up = "Follow-up / 后续说明: If anything is still unclear, it can be discussed and confirmed in the later review meeting."

    add_title(doc, title)
    add_intro(doc, purpose, instructions)

    doc.add_heading("Part A. Group Overview / 小组概况", level=1)
    add_part_a_table(doc)

    doc.add_paragraph()
    doc.add_heading("Part B. Requirement Items / 需求明细", level=1)
    add_part_b_table(doc)

    doc.add_paragraph()
    add_hint_list(doc, group["zh_hints"] if language == "zh" else group["en_hints"], hint_heading)

    doc.add_paragraph()
    doc.add_paragraph(follow_up)
    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for group in GROUPS:
        zh_doc = build_document("zh", group)
        zh_doc.save(OUTPUT_DIR / group["zh_filename"])

        en_doc = build_document("en", group)
        en_doc.save(OUTPUT_DIR / group["en_filename"])

    print("PASS: generated 8 requirement form templates")


if __name__ == "__main__":
    main()
